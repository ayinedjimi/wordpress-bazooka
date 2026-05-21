"""DOCX report exporter — Professional Word document with findings, remediation, and compliance."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

if TYPE_CHECKING:
    from core.engine import ScanContext

SEVERITY_COLORS = {
    "CRITICAL": RGBColor(241, 40, 40),
    "HIGH": RGBColor(255, 106, 0),
    "MEDIUM": RGBColor(232, 163, 23),
    "LOW": RGBColor(255, 216, 0),
    "INFO": RGBColor(15, 130, 255),
}


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _clean_xml(text: str) -> str:
    """Remove characters that are invalid in XML (control chars except tab/newline/cr)."""
    import re
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', str(text))


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_clean_xml(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _severity_text(severity: str) -> str:
    icons = {"CRITICAL": "[!!]", "HIGH": "[!]", "MEDIUM": "[*]", "LOW": "[-]", "INFO": "[i]"}
    return f"{icons.get(severity, '')} {severity}"


def generate_docx_report(ctx: ScanContext, output_dir: Path) -> Path:
    """Generate a professional DOCX report."""
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # -- Cover page --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\nRAPPORT D'AUDIT DE SECURITE\nWORDPRESS\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(88, 81, 219)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{ctx.target.domain}\n")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(26, 23, 42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = ctx.target.meta
    date_str = datetime.utcnow().strftime("%d/%m/%Y")
    run = p.add_run(f"\nDate: {date_str}\nProfil: {meta.profile}\nRequetes: {meta.total_requests}\n")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nCONFIDENTIEL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(241, 40, 40)

    doc.add_page_break()

    # -- Disclaimer --
    _add_para(doc,
        "Ce rapport a ete produit dans le cadre d'un audit de securite autorise. "
        "Les tests ont ete realises conformement au perimetre defini. "
        "L'utilisation de cet outil en dehors d'un cadre legal autorise est strictement interdite.",
        italic=True, size=9,
    )

    if meta.authorization_ref:
        _add_para(doc, f"Reference d'autorisation: {meta.authorization_ref}", bold=True)

    doc.add_paragraph()

    # -- Table of contents placeholder --
    _add_heading(doc, "Table des matieres", 1)
    _add_para(doc, "1. Resume executif\n2. Synthese des vulnerabilites\n3. Findings detailles\n"
                    "4. Plan de remediation\n5. Donnees de reconnaissance\n6. Annexes")
    doc.add_page_break()

    # -- 1. Executive Summary --
    _add_heading(doc, "1. Resume executif", 1)

    counts = ctx.severity_counts
    max_cvss = ctx.max_cvss
    total = len(ctx.findings)

    _add_para(doc, f"Score de risque maximal: {max_cvss}/10", bold=True, size=14)

    # Severity summary table
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[1].cells[i].text = str(counts.get(h, 0))
        for paragraph in table.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        for paragraph in table.rows[1].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(14)
                run.bold = True

    doc.add_paragraph()

    # Target info
    info_parts = []
    if ctx.target.wp_version:
        info_parts.append(f"WordPress {ctx.target.wp_version}")
    if ctx.target.waf_detected:
        info_parts.append(f"WAF: {ctx.target.waf_detected}")
    if ctx.target.users:
        info_parts.append(f"Utilisateurs: {len(ctx.target.users)}")
    if ctx.target.plugins:
        info_parts.append(f"Plugins: {len(ctx.target.plugins)}")
    if info_parts:
        _add_para(doc, " | ".join(info_parts), size=11)

    doc.add_page_break()

    # -- 2. Findings Summary Table --
    _add_heading(doc, "2. Synthese des vulnerabilites", 1)

    findings = sorted(ctx.findings, key=lambda f: f.cvss_score, reverse=True)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Severite", "CVSS", "Titre", "Module"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    for f in findings:
        row = table.add_row().cells
        row[0].text = f.id
        row[1].text = f.severity.value
        row[2].text = str(f.cvss_score)
        row[3].text = f.title[:80]
        row[4].text = f.module
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    doc.add_page_break()

    # -- 3. Detailed Findings --
    _add_heading(doc, "3. Findings detailles", 1)

    for f in findings:
        sev = f.severity.value
        color = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))

        p = doc.add_paragraph()
        run = p.add_run(f"{_severity_text(sev)} ")
        run.bold = True
        run.font.color.rgb = color
        run.font.size = Pt(10)
        run = p.add_run(_clean_xml(f.title))
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        run = p.add_run(f"CVSS: {f.cvss_score}  |  Confiance: {f.confidence.value}  |  Module: {f.module}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

        if f.description:
            _add_para(doc, _clean_xml(f.description), size=9)

        if f.evidence and f.evidence.request:
            p = doc.add_paragraph()
            run = p.add_run("Evidence: ")
            run.bold = True
            run.font.size = Pt(8)
            run = p.add_run(_clean_xml(f.evidence.request[:300]))
            run.font.size = Pt(8)
            run.font.name = "Consolas"

        if f.impact:
            p = doc.add_paragraph()
            run = p.add_run("Impact: ")
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(_clean_xml(f.impact))
            run.font.size = Pt(9)

        if f.remediation:
            p = doc.add_paragraph()
            run = p.add_run("Remediation: ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 128, 0)
            run = p.add_run(_clean_xml(f.remediation))
            run.font.size = Pt(9)

        # Compliance tags
        tags = []
        if f.compliance.owasp_2021:
            tags.append(f"OWASP: {f.compliance.owasp_2021}")
        if f.compliance.cwe:
            tags.append(f.compliance.cwe)
        if f.compliance.mitre_attack:
            tags.append(f"MITRE: {f.compliance.mitre_attack}")
        if tags:
            _add_para(doc, " | ".join(tags), size=8, italic=True)

        doc.add_paragraph()  # Spacer

    doc.add_page_break()

    # -- 4. Remediation Plan --
    _add_heading(doc, "4. Plan de remediation", 1)
    _add_para(doc, "Actions classees par priorite (CRITICAL en premier).", italic=True)

    severity_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3, "INFO": 4}
    remed_map: dict[str, dict] = {}
    for f in findings:
        key = f.remediation or "Aucune remediation specifiee"
        if key not in remed_map:
            remed_map[key] = {"action": key, "severity": f.severity.value, "rank": severity_order.get(f.severity.value, 5), "ids": []}
        remed_map[key]["ids"].append(f.id)
        if severity_order.get(f.severity.value, 5) < remed_map[key]["rank"]:
            remed_map[key]["severity"] = f.severity.value
            remed_map[key]["rank"] = severity_order.get(f.severity.value, 5)

    plan = sorted(remed_map.values(), key=lambda x: x["rank"])

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Priorite"
    hdr[1].text = "Action"
    hdr[2].text = "Findings"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    for item in plan:
        row = table.add_row().cells
        row[0].text = item["severity"]
        row[1].text = item["action"][:150]
        row[2].text = ", ".join(item["ids"][:5])
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)

    doc.add_page_break()

    # -- 5. Recon Data --
    _add_heading(doc, "5. Donnees de reconnaissance", 1)

    if ctx.target.users:
        _add_heading(doc, "Utilisateurs detectes", 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["ID", "Username", "Display Name", "Email"]):
            table.rows[0].cells[i].text = h
        for u in ctx.target.users:
            row = table.add_row().cells
            row[0].text = str(u.id)
            row[1].text = u.username
            row[2].text = u.display_name
            row[3].text = u.email or "-"

    if ctx.target.plugins:
        _add_heading(doc, "Plugins detectes", 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Slug", "Version", "CVEs"]):
            table.rows[0].cells[i].text = h
        for p in ctx.target.plugins:
            row = table.add_row().cells
            row[0].text = p.slug
            row[1].text = p.version or "?"
            row[2].text = str(len(p.cves))

    # -- Footer --
    doc.add_page_break()
    _add_heading(doc, "6. Annexes", 1)
    _add_para(doc, f"WordPress BAZOOKA v{meta.bazooka_version}", bold=True)
    _add_para(doc, f"Genere le {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    _add_para(doc, f"Total: {total} vulnerabilites | {meta.total_requests} requetes | Profil: {meta.profile}")
    _add_para(doc, "Ce rapport a ete genere automatiquement et ne remplace pas l'expertise humaine.", italic=True, size=9)

    # Save
    output_file = output_dir / f"RAPPORT_BAZOOKA_{ctx.target.domain}.docx"
    doc.save(str(output_file))
    return output_file
